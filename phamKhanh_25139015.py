#name: Phạm Duy Khánh
#MSSV: 25139015
from docx import Document
import re

def number_questions(input_path, output_path):
    # Mở file Word đầu vào
    doc = Document(input_path)
    
    # Bước 1: Xóa số thứ tự cũ (ví dụ: "câu hỏi 1" -> "câu hỏi")
    for para in doc.paragraphs:
        para.text = re.sub(r'câu hỏi \d{1,3}', 'câu hỏi', para.text)
    
    # Duyệt qua bảng nếu có
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = re.sub(r'câu hỏi \d{1,3}', 'câu hỏi', cell.text)
    
    # Bước 2: Đánh số mới
    search_text = "câu hỏi"
    i = 1
    
    # Duyệt qua đoạn văn
    for para in doc.paragraphs:
        if search_text in para.text:
            new_text = para.text
            while search_text in new_text:
                new_text = new_text.replace(search_text, f"câu hỏi {i}", 1)
                i += 1
            para.text = new_text
    
    # Duyệt qua bảng
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if search_text in cell.text:
                    new_text = cell.text
                    while search_text in new_text:
                        new_text = new_text.replace(search_text, f"câu hỏi {i}", 1)
                        i += 1
                    cell.text = new_text
    
    # Lưu file mới
    doc.save(output_path)
    print(f"Đã đánh số {i-1} câu hỏi và lưu vào file: {output_path}")
    if i - 1 != 159:
        print("Cảnh báo: Số lượng 'câu hỏi' tìm thấy không phải 159. Kiểm tra lại file.")

# Đường dẫn chính xác cho file đầu vào và đầu ra
input_file = r"C:\Users\Admin\Documents\forrm\FileGoc.docx"
output_file = r"C:\Users\Admin\Documents\forrm\FileKetQua.docx"
number_questions(input_file, output_file)
# File: phamKhanh_25139015.py
# Họ tên: Phạm Khánh
# MSSV: 25139015

def display_info():
    info = {
        "Họ tên": "Phạm Khánh",
        "MSSV": "25139015",
        "Email": "phamkhanh@edu.vn"  # Thêm email nếu có
    }
    for key, value in info.items():
        print(f"{key}: {value}")

if __name__ == "__main__":
    print("Thông tin cá nhân:")
    display_info()
